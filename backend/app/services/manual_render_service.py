"""
DOCX-Handbuch-Renderer fuer Storyboards.
"""
from __future__ import annotations

import json
import re
import time
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import datetime
from datetime import date
from pathlib import Path
from typing import Any, Callable

from app.config import settings
from app.models import AiProvider, Scene, StoryboardJson, TextPanel
from app.services.storyboard_service import StoryboardService

_VALID_ADDRESS = {"du", "sie", "neutral"}
_VALID_WRITING = {"sachlich", "leicht_verstaendlich", "technisch_detailliert"}
_VALID_DETAIL = {"kurz", "standard", "ausfuehrlich"}

_ADDRESS_CLAUSES = {
    "du": "Verwende durchgehend die Du-Form (direkte Ansprache, 'du', 'dein').",
    "sie": "Verwende durchgehend die Sie-Form (formelle Ansprache, 'Sie', 'Ihr').",
    "neutral": (
        "Schreibe unpersönlich ohne direkte Ansprache (kein 'du', kein 'Sie'). "
        "Verwende stattdessen Infinitivkonstruktionen oder 'Der Benutzer ...' / 'Es wird ...'."
    ),
}
_STYLE_CLAUSES = {
    "sachlich": (
        "Schreibstil: sachlich und präzise. Vollständige Sätze, klare Struktur, keine Umgangssprache, "
        "keine unnötigen Ausschmückungen."
    ),
    "leicht_verstaendlich": (
        "Schreibstil: leicht verständlich. Kurze Sätze, einfache Wörter, aktive Formulierungen. "
        "Fachbegriffe nur wenn nötig, dann kurz erläutert."
    ),
    "technisch_detailliert": (
        "Schreibstil: technisch detailliert. Fachterminologie konsequent verwenden, "
        "Abläufe präzise und vollständig beschreiben."
    ),
}
_DETAIL_CLAUSES = {
    "kurz": (
        "Detailtiefe: kurz und kompakt. Maximal 2–3 Sätze pro Segment. "
        "Formuliere prägnant, lass weniger wichtige Details weg."
    ),
    "standard": (
        "Detailtiefe: ausgewogen. 3–5 Sätze pro Segment. "
        "Erkläre den Schritt vollständig, ohne zu überladen."
    ),
    "ausfuehrlich": (
        "Detailtiefe: ausführlich. 5–8 Sätze pro Segment. "
        "Erkläre den gezeigten Schritt umfassend: was zu tun ist, warum, "
        "was dabei zu beachten ist und was das Ergebnis sein soll. "
        "Nutze dabei ausschließlich Informationen, die im Storyboard vorhanden sind — "
        "erfinde keine neuen technischen Fakten."
    ),
}

_MAX_IMAGES_PER_KI_CALL = 4  # Obergrenze Bilder pro KI-Aufruf (Provider-agnostisch)


@dataclass(frozen=True)
class ManualOptimizationOptions:
    lang: str
    address_style: str = "sie"
    writing_style: str = "sachlich"
    detail_level: str = "standard"
    doc_title: str = ""


@dataclass
class _HandbuchSessionEntry:
    scene_id: str
    heading: str
    segments: list[dict[str, str]]   # [{image, text}, ...]


@dataclass
class _HandbuchSession:
    """In-Memory-Session für einen Handbuch-Render-Vorgang (nicht in session_store)."""
    video_id: str
    lang: str
    options: ManualOptimizationOptions
    source_name: str
    total_scenes: int
    completed: list[_HandbuchSessionEntry] = field(default_factory=list)
    doc_title: str = ""

    def context_summary(self) -> str:
        """Kontext der zuletzt fertiggestellten Szenen — vollständiger Text für Stil-Kontinuität."""
        lines = [
            f"Handbuch-Session: {self.source_name}, Sprache {self.lang}, "
            f"{len(self.completed)}/{self.total_scenes} Szenen fertig.",
            f"Stil: {self.options.address_style} | {self.options.writing_style} | {self.options.detail_level}",
            "Zuletzt fertiggestellte Szene(n) — halte Stil und Übergangsformulierungen konsistent:",
        ]
        # Die letzten 3 Szenen mit vollständigen Segment-Texten mitgeben
        for entry in self.completed[-3:]:
            lines.append(f"  Szene {entry.scene_id!r}: {entry.heading!r}")
            for seg in entry.segments:
                lines.append(f"    [{seg['image']}]: {seg['text']}")
        return "\n".join(lines)


class ManualRenderService:
    """Erzeugt bearbeitbare DOCX-Handbuecher aus bestehenden Storyboards."""

    def __init__(self) -> None:
        self._storyboard_svc = StoryboardService()

    def render_manual(
        self,
        video_id: str,
        lang: str,
        optimize: bool = False,
        ai_provider: AiProvider | None = None,
        ai_model: str | None = None,
        debug_callback: Callable[[str, str], None] | None = None,
        address_style: str = "sie",
        writing_style: str = "sachlich",
        detail_level: str = "standard",
    ) -> Path:
        if address_style not in _VALID_ADDRESS:
            raise ValueError(
                f"address_style muss einer der Werte sein: {', '.join(sorted(_VALID_ADDRESS))}"
            )
        if writing_style not in _VALID_WRITING:
            raise ValueError(
                f"writing_style muss einer der Werte sein: {', '.join(sorted(_VALID_WRITING))}"
            )
        if detail_level not in _VALID_DETAIL:
            raise ValueError(
                f"detail_level muss einer der Werte sein: {', '.join(sorted(_VALID_DETAIL))}"
            )

        options = ManualOptimizationOptions(
            lang=lang,
            address_style=address_style,
            writing_style=writing_style,
            detail_level=detail_level,
        )

        storyboard = self._storyboard_svc.load(video_id)
        manual_storyboard = deepcopy(storyboard)
        doc_title = ""
        ai_info = ""
        if optimize:
            provider_name = ai_provider.value if ai_provider else settings.ai_provider
            ai_info = f"{provider_name} / {ai_model}" if ai_model else provider_name
            manual_storyboard, doc_title = self._optimize_for_manual(
                storyboard,
                options,
                ai_provider=ai_provider,
                ai_model=ai_model,
                debug_callback=debug_callback,
            )
            self._save_manual_storyboard(video_id, lang, manual_storyboard)

        output_dir = settings.render_output_dir / video_id
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = self._build_filename(storyboard, lang, doc_title)
        out_path = output_dir / filename
        self._write_docx(manual_storyboard, lang, out_path, ai_info=ai_info)
        return out_path

    def _get_provider(self, provider: AiProvider | None, model: str | None):
        provider_name = provider.value if provider else settings.ai_provider
        if provider_name == AiProvider.GEMINI.value:
            from app.services.gemini_provider import GeminiProvider
            return GeminiProvider(model=model)
        if provider_name == AiProvider.OPENAI.value:
            from app.services.openai_provider import OpenAiProvider
            return OpenAiProvider(model=model)
        if provider_name == AiProvider.AZURE_OPENAI.value:
            from app.services.azure_openai_provider import AzureOpenAiProvider
            return AzureOpenAiProvider(model=model)
        if provider_name == AiProvider.AZURE_COGNITIVE.value:
            from app.services.azure_cognitive_provider import AzureCognitiveProvider
            return AzureCognitiveProvider(model=model)
        raise ValueError(f"Unbekannter AI-Provider: {provider_name}")

    def _optimize_for_manual(
        self,
        storyboard: StoryboardJson,
        options: ManualOptimizationOptions,
        ai_provider: AiProvider | None,
        ai_model: str | None,
        debug_callback: Callable[[str, str], None] | None = None,
    ) -> tuple[StoryboardJson, str]:
        provider = self._get_provider(ai_provider, ai_model)
        source_name = Path(storyboard.source_video).stem if storyboard.source_video else storyboard.video_id
        total_scenes = len(storyboard.scenes)
        total_images = sum(len(s.image_group or []) for s in storyboard.scenes)

        if debug_callback:
            debug_callback("manual-options", json.dumps({
                "lang": options.lang,
                "address_style": options.address_style,
                "writing_style": options.writing_style,
                "detail_level": options.detail_level,
                "scene_count": total_scenes,
                "image_count": total_images,
                "mode": "scene-by-scene",
            }, ensure_ascii=False))

        # Neue In-Memory-Session für diesen Handbuch-Vorgang
        session = _HandbuchSession(
            video_id=storyboard.video_id,
            lang=options.lang,
            options=options,
            source_name=source_name,
            total_scenes=total_scenes,
        )

        optimized = deepcopy(storyboard)

        for idx, source_scene in enumerate(storyboard.scenes):
            scene_num = idx + 1
            source_lang_text = source_scene.texts.get(options.lang, TextPanel())
            input_had_text = bool(
                (source_lang_text.body or "").strip()
                or (source_lang_text.speaker_notes or "").strip()
            )

            chunks = self._chunk_image_group(source_scene.image_group, _MAX_IMAGES_PER_KI_CALL)
            chunk_total = len(chunks)

            if debug_callback:
                debug_callback(
                    "manual-status",
                    f"Szene {scene_num}/{total_scenes}: {source_scene.scene_id} "
                    f"({len(source_scene.image_group)} Bild(er), {chunk_total} Chunk(s))…"
                )

            all_panels: list[TextPanel] = []
            all_session_segments: list[dict[str, str]] = []

            for chunk_idx, chunk_images in enumerate(chunks):
                prompt = self._build_scene_prompt(
                    source_scene, options, session, scene_num,
                    chunk_images=chunk_images,
                    chunk_idx=chunk_idx,
                    chunk_total=chunk_total,
                    prior_segments=all_session_segments,
                )

                if debug_callback:
                    debug_callback(f"manual-prompt-{scene_num}-chunk-{chunk_idx}", prompt)

                frame_paths = [
                    settings.frames_dir / storyboard.video_id / fname
                    for fname in chunk_images
                    if (settings.frames_dir / storyboard.video_id / fname).exists()
                ]

                raw = self._call_provider(provider, prompt, frame_paths, scene_num, chunk_idx, debug_callback)

                if debug_callback:
                    debug_callback(f"manual-response-{scene_num}-chunk-{chunk_idx}", raw[:3000])

                data = self._parse_json(raw)

                # Titel nur aus dem ersten Chunk der ersten Szene
                if not session.doc_title and chunk_idx == 0:
                    session.doc_title = str(data.get("title", "")).strip()

                # Segmente validieren — erwartet genau len(chunk_images) Einträge
                segments_raw = data.get("segments")
                if not isinstance(segments_raw, list):
                    scenes_list = data.get("scenes")
                    if isinstance(scenes_list, list) and len(scenes_list) == 1:
                        segments_raw = scenes_list[0].get("segments")
                if not isinstance(segments_raw, list):
                    preview = raw[:300].replace("\n", " ")
                    raise ValueError(
                        f"Handbuch-KI-Antwort Szene {scene_num} Chunk {chunk_idx + 1} "
                        f"enthält keine segments-Liste. Antwort-Anfang: {preview}"
                    )

                expected_count = len(chunk_images)
                if len(segments_raw) != expected_count:
                    raise ValueError(
                        f"Handbuch-KI-Antwort Szene {scene_num} Chunk {chunk_idx + 1}: "
                        f"Segmentanzahl stimmt nicht (erwartet {expected_count}, erhalten {len(segments_raw)})."
                    )

                for seg_idx, segment in enumerate(segments_raw):
                    if not isinstance(segment, dict):
                        raise ValueError(
                            f"Handbuch-KI-Antwort Szene {scene_num} Chunk {chunk_idx + 1}, "
                            f"Segment {seg_idx + 1}: kein Objekt."
                        )
                    expected_image = chunk_images[seg_idx]
                    segment["image"] = expected_image
                    seg_text = segment.get("text", "")
                    if not isinstance(seg_text, str):
                        raise ValueError(
                            f"Handbuch-KI-Antwort Szene {scene_num} Chunk {chunk_idx + 1}, "
                            f"Segment {seg_idx + 1}: 'text' ist kein String."
                        )
                    if input_had_text and not seg_text.strip():
                        raise ValueError(
                            f"Handbuch-KI-Antwort Szene {scene_num} Chunk {chunk_idx + 1}, "
                            f"Segment {seg_idx + 1}: 'text' ist leer, obwohl die Eingabe Text enthielt."
                        )
                    all_panels.append(TextPanel(
                        heading=source_lang_text.heading,
                        body=seg_text,
                        speaker_notes=seg_text,
                    ))
                    all_session_segments.append({"image": expected_image, "text": seg_text})

            # Alle Chunks dieser Szene zusammengeführt → slide_panels schreiben
            optimized.scenes[idx].slide_panels[options.lang] = all_panels

            session.completed.append(_HandbuchSessionEntry(
                scene_id=source_scene.scene_id,
                heading=source_lang_text.heading or source_scene.scene_id,
                segments=all_session_segments,
            ))

            if debug_callback:
                debug_callback(
                    "manual-progress",
                    f"Szene {scene_num}/{total_scenes} fertig ✓ "
                    f"({len(all_panels)} Segment(e) in {chunk_total} Chunk(s) → slide_panels[{options.lang}])"
                )

        optimized.metadata = {
            **optimized.metadata,
            "manual_optimization": {
                "language": options.lang,
                "address_style": options.address_style,
                "writing_style": options.writing_style,
                "detail_level": options.detail_level,
                "mode": "scene-by-scene",
                "scene_count": total_scenes,
                "instruction": (
                    "KI-Antwort wurde szenenweise genutzt, um Sprech-Skript-Text in Handbuch-Prosa umzuwandeln. "
                    "Szenenstruktur und Bildreferenzen stammen unveraendert aus dem Storyboard."
                ),
            },
        }
        return optimized, session.doc_title

    def _build_manual_prompt(self, storyboard: StoryboardJson, options: ManualOptimizationOptions) -> str:
        scenes_payload: list[dict[str, Any]] = []
        for scene in storyboard.scenes:
            text = scene.texts.get(options.lang, TextPanel())
            slide_panels = scene.slide_panels.get(options.lang, []) if scene.slide_panels else []
            scenes_payload.append({
                "scene_id": scene.scene_id,
                "image_group": scene.image_group,
                "text": text.model_dump(),
                "slide_panels": [
                    panel.model_dump() if hasattr(panel, "model_dump") else panel
                    for panel in slide_panels
                ],
            })

        source_name = Path(storyboard.source_video).stem if storyboard.source_video else ""
        framing = (
            "Du bist ein technischer Redakteur und erstellst ein Benutzerhandbuch aus einem Video-Storyboard.\n"
            "Das Storyboard stammt aus einer Videoaufzeichnung und enthält Sprech-Skript-Texte.\n"
            "Deine Aufgabe ist es, für jedes Bild einen einzigen, zusammenhängenden Handbuch-Fließtext zu verfassen.\n"
            "Der Text beschreibt, was auf dem Bild zu sehen ist und erklärt den gezeigten Schritt vollständig. "
            "Er ersetzt sowohl die Bild-Erklärung als auch den Anleitungstext — alles in einem Absatz.\n\n"
            f"Quelldatei: {source_name}"
        )
        address_clause = _ADDRESS_CLAUSES[options.address_style]
        style_clause = _STYLE_CLAUSES[options.writing_style]
        detail_clause = _DETAIL_CLAUSES[options.detail_level]
        constraints = (
            "Erlaubt:\n"
            "- Skript-Text sprachlich in Handbuch-Prosa umformulieren\n"
            "- Satzbau, Tempo und Stil an das Handbuch-Format anpassen\n"
            "- Wiederholungen reduzieren und vorhandene Informationen klarer strukturieren\n"
            "- Den vorhandenen Inhalt je nach Detailtiefe ausführlicher erklären\n"
            "- Kleine Übergänge für besseren Lesefluss\n\n"
            "Verboten:\n"
            "- Neue technische Fakten, Werte, Ports, Namen, Warnungen oder Systemkomponenten erfinden\n"
            "- Neue Schritte oder Handlungsanweisungen hinzufügen, die im Original nicht vorkommen\n"
            "- Szenen erzeugen, löschen, zusammenführen, aufteilen oder verschieben\n"
            "- scene_id oder image_group verändern\n"
            "- Bildreihenfolge verändern\n"
            "- Halluzinationen jeglicher Art\n\n"
            "Videosprache, die NICHT im Handbuch erscheinen darf:\n"
            "- 'Hallo zusammen', 'hier sehen wir', 'nun schauen wir', 'wie im Video gezeigt',\n"
            "  'Bild X zeigt uns', 'wir klicken jetzt', 'jetzt sehen Sie', 'schauen wir uns an'"
        )
        format_spec = (
            "Gib ausschließlich valides JSON zurück.\n\n"
            "Rückgabeformat:\n"
            '{"title":"Kurzer prägnanter Handbuch-Titel (max. 6 Wörter, dateinamensicher, keine Sonderzeichen außer Bindestrich)",'
            '"scenes":[{"scene_id":"scene_001","image_group":["frame_001.jpg"],'
            '"segments":[{"image":"frame_001.jpg",'
            '"text":"Vollständiger Handbuch-Fließtext für dieses Bild als ein zusammenhängender Absatz"}]}]}\n\n'
            "JSON-Regeln:\n"
            "- 'title': kurzer, beschreibender Dokumenttitel, max. 6 Wörter, nur Buchstaben, Ziffern und Bindestriche, keine Umlaute\n"
            "- scenes-Liste: exakt dieselbe Anzahl Einträge wie die Eingabe\n"
            "- scene_id und image_group: exakt identisch zur Eingabe\n"
            "- segments: exakt so viele Einträge wie Bilder in image_group\n"
            "- Jedes Segment enthält: 'image' (identischer Dateiname aus image_group), 'text' (nicht-leerer String)\n"
            "- 'text' darf nicht leer sein, wenn die Eingabe Text enthielt"
        )
        payload = (
            f"Sprache: {options.lang}\n"
            f"Storyboard:\n{json.dumps({'scenes': scenes_payload}, ensure_ascii=False)}"
        )
        return "\n\n".join([framing, address_clause, style_clause, detail_clause, constraints, format_spec, payload])

    @staticmethod
    def _chunk_image_group(image_group: list[str], chunk_size: int) -> list[list[str]]:
        """Teilt image_group in Chunks von max. chunk_size Einträgen."""
        return [image_group[i: i + chunk_size] for i in range(0, len(image_group), chunk_size)]

    def _call_provider(
        self,
        provider: Any,
        prompt: str,
        frame_paths: list[Path],
        scene_num: int,
        chunk_idx: int,
        debug_callback: Callable[[str, str], None] | None,
    ) -> str:
        """KI-Aufruf mit Retry bei transienten Provider-Fehlern (503/429)."""
        last_exc: Exception | None = None
        for attempt in range(3):
            try:
                if frame_paths:
                    return provider.complete_text_with_images(prompt, frame_paths)
                return provider.complete_text(prompt)
            except Exception as exc:
                last_exc = exc
                msg = str(exc)
                if any(code in msg for code in ("503", "429", "UNAVAILABLE", "overloaded")):
                    wait = 15 * (attempt + 1)
                    if debug_callback:
                        debug_callback(
                            "manual-status",
                            f"Szene {scene_num} Chunk {chunk_idx + 1}: Provider-Fehler ({msg[:80]}), "
                            f"Retry in {wait}s…"
                        )
                    time.sleep(wait)
                else:
                    raise
        raise last_exc  # type: ignore[misc]

    def _build_scene_prompt(
        self,
        scene: Scene,
        options: ManualOptimizationOptions,
        session: "_HandbuchSession",
        scene_num: int,
        chunk_images: list[str],
        chunk_idx: int = 0,
        chunk_total: int = 1,
        prior_segments: list[dict[str, str]] | None = None,
    ) -> str:
        is_first_chunk = chunk_idx == 0
        is_chunked = chunk_total > 1

        if is_chunked:
            framing = (
                f"Du bist ein technischer Redakteur und erstellst "
                f"Teil {chunk_idx + 1} von {chunk_total} der Szene {scene_num} "
                f"(von insgesamt {session.total_scenes} Szenen) eines Benutzerhandbuchs.\n"
                f"Quelldatei: {session.source_name}\n"
                f"Dieser Teilblock enthält {len(chunk_images)} Bild(er).\n"
                "Dir werden die Bildschirmaufnahmen dieses Blocks als Bilder mitgeschickt. "
                "Nutze den visuellen Inhalt als Grundlage, ergänzt durch die Storyboard-Texte. "
                "Formuliere jeden Bildabschnitt als zusammenhängenden Handbuch-Fließtext (ein Absatz pro Bild)."
            )
        else:
            framing = (
                f"Du bist ein technischer Redakteur und erstellst Szene {scene_num} von {session.total_scenes} "
                f"eines Benutzerhandbuchs aus einem Video-Storyboard.\n"
                f"Quelldatei: {session.source_name}\n"
                "Dir werden die Bildschirmaufnahmen der Szene als Bilder mitgeschickt. "
                "Nutze den visuellen Inhalt als Grundlage, ergänzt durch die Storyboard-Texte. "
                "Formuliere jeden Bildabschnitt als zusammenhängenden Handbuch-Fließtext (ein Absatz pro Bild)."
            )

        address_clause = _ADDRESS_CLAUSES[options.address_style]
        style_clause = _STYLE_CLAUSES[options.writing_style]
        detail_clause = _DETAIL_CLAUSES[options.detail_level]

        constraints = (
            "Erlaubt: Skript-Text in Handbuch-Prosa umformulieren, Stil anpassen, Wiederholungen reduzieren.\n"
            "Verboten: neue technische Fakten erfinden, Szenen/Bilder verändern, Halluzinationen.\n"
            "Kein Videojargon: 'hier sehen wir', 'wir klicken jetzt', 'schauen wir uns an' usw."
        )

        # Session-Kontext: letzte fertige Szenen als Stil-Anker
        context_parts: list[str] = []
        if session.completed:
            context_parts.append(
                f"Session-Kontext (bereits fertiggestellte Szenen):\n{session.context_summary()}"
            )

        # Intra-Szene-Kontext: bereits fertige Segmente früherer Chunks derselben Szene
        if prior_segments and not is_first_chunk:
            prior_lines = [
                f"  [{seg['image']}]: {seg['text'][:200]}{'…' if len(seg['text']) > 200 else ''}"
                for seg in prior_segments
            ]
            context_parts.append(
                "Bereits fertiggestellte Segmente dieser Szene (Stil-Anker — nahtlos fortsetzen):\n"
                + "\n".join(prior_lines)
            )

        context_block = "\n\n".join(context_parts)

        # Bildliste nur für diesen Chunk
        image_list = "\n".join(
            f"  Bild {i + 1} von {len(chunk_images)}: {fname}"
            for i, fname in enumerate(chunk_images)
        )
        first_frame = chunk_images[0] if chunk_images else "frame_001.jpg"

        # Titel nur im ersten Chunk der ersten Szene anfordern
        want_title = is_first_chunk and scene_num == 1
        if want_title:
            format_spec = (
                "Gib ausschließlich valides JSON zurück. "
                "Die Reihenfolge der Bilder ist EXAKT einzuhalten:\n"
                f'{{"title":"...","segments":[{{"image":"{first_frame}","text":"Handbuch-Fließtext..."}},...]}}\n'
                '"title": Kurzer Handbuch-Titel (max. 6 Wörter, nur Buchstaben/Ziffern/Bindestriche, keine Umlaute)\n'
                f'"segments": exakt {len(chunk_images)} Einträge in der Reihenfolge:\n'
                f"{image_list}\n"
                'Je Eintrag: "image" = exakter Dateiname aus obiger Liste, "text" = Handbuch-Fließtext (nicht leer).'
            )
        else:
            format_spec = (
                "Gib ausschließlich valides JSON zurück. "
                "Die Reihenfolge der Bilder ist EXAKT einzuhalten:\n"
                f'{{"segments":[{{"image":"{first_frame}","text":"Handbuch-Fließtext..."}},...]}}\n'
                f'"segments": exakt {len(chunk_images)} Einträge in der Reihenfolge:\n'
                f"{image_list}\n"
                'Je Eintrag: "image" = exakter Dateiname aus obiger Liste, "text" = Handbuch-Fließtext (nicht leer).'
            )

        scene_text = scene.texts.get(options.lang, TextPanel())
        text_payload = {
            "scene_id": scene.scene_id,
            "heading": scene_text.heading,
            "body": scene_text.body,
            "speaker_notes": scene_text.speaker_notes,
            "images": chunk_images,
        }
        payload = (
            f"Sprache: {options.lang}\n"
            f"Diese Szene:\n{json.dumps(text_payload, ensure_ascii=False)}"
        )

        parts = [framing, address_clause, style_clause, detail_clause, constraints]
        if context_block:
            parts.append(context_block)
        parts += [format_spec, payload]
        return "\n\n".join(parts)

    def _parse_json(self, raw: str) -> dict[str, Any]:
        cleaned = raw.strip()
        match = re.search(r"```(?:json)?\s*([\s\S]+?)\s*```", cleaned)
        if match:
            cleaned = match.group(1).strip()
        start = cleaned.find("{")
        end = cleaned.rfind("}") + 1
        if start >= 0 and end > start:
            cleaned = cleaned[start:end]

        # Erster Versuch: direkt parsen
        try:
            data = json.loads(cleaned)
            if not isinstance(data, dict):
                raise ValueError("Handbuch-KI-Antwort ist kein JSON-Objekt.")
            return data
        except json.JSONDecodeError:
            pass

        # Zweiter Versuch: ungültige Backslash-Escapes innerhalb von Strings reparieren.
        # Gemini schreibt manchmal \E, \e, \Z o.ä. als Literal in Texten statt \\ zu verwenden.
        # Strategie: alle \X ersetzen wo X kein gültiger JSON-Escape-Zeichentyp ist.
        _VALID_ESCAPES = set('"\\/ bfnrtu')

        def fix_invalid_escapes(s: str) -> str:
            result: list[str] = []
            i = 0
            in_string = False
            while i < len(s):
                ch = s[i]
                if ch == '"' and (i == 0 or s[i - 1] != '\\'):
                    in_string = not in_string
                    result.append(ch)
                    i += 1
                elif in_string and ch == '\\' and i + 1 < len(s):
                    next_ch = s[i + 1]
                    if next_ch in _VALID_ESCAPES:
                        result.append(ch)
                        result.append(next_ch)
                        i += 2
                    else:
                        # Ungültiger Escape: Backslash verdoppeln
                        result.append('\\\\')
                        result.append(next_ch)
                        i += 2
                else:
                    result.append(ch)
                    i += 1
            return ''.join(result)

        repaired = fix_invalid_escapes(cleaned)
        try:
            data = json.loads(repaired)
            if not isinstance(data, dict):
                raise ValueError("Handbuch-KI-Antwort ist kein JSON-Objekt.")
            return data
        except json.JSONDecodeError as exc:
            preview = raw[:600].replace("\n", " ")
            raise ValueError(
                f"Handbuch-KI-Antwort ist kein valides JSON (nach Escape-Repair). "
                f"Länge: {len(raw)} Zeichen. Antwort-Anfang: {preview}"
            ) from exc

    def _save_manual_storyboard(
        self,
        video_id: str,
        lang: str,
        storyboard: StoryboardJson,
    ) -> Path:
        out_dir = settings.ai_output_dir / video_id
        out_dir.mkdir(parents=True, exist_ok=True)
        path = out_dir / f"manual_storyboard_{lang}.json"
        path.write_text(storyboard.model_dump_json(indent=2), encoding="utf-8")
        return path

    def _build_filename(self, storyboard: StoryboardJson, lang: str, doc_title: str) -> str:
        if doc_title:
            # KI-generierter Titel: bereinigen auf dateinamensichere Zeichen
            safe = re.sub(r"[^\w\-]", "-", doc_title, flags=re.UNICODE)
            safe = re.sub(r"-{2,}", "-", safe).strip("-")
        else:
            # Fallback: Quelldateiname ohne Erweiterung
            safe = Path(storyboard.source_video).stem if storyboard.source_video else "handbuch"
            safe = re.sub(r"[^\w\-]", "-", safe, flags=re.UNICODE)
            safe = re.sub(r"-{2,}", "-", safe).strip("-")
        return f"{safe}_{lang}.docx"

    def _write_docx(self, storyboard: StoryboardJson, lang: str, out_path: Path, ai_info: str = "") -> None:
        try:
            from docx import Document
            from docx.enum.section import WD_ORIENT
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_BREAK
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Mm, Pt
        except ImportError as exc:
            raise RuntimeError("python-docx ist nicht installiert. Bitte requirements installieren.") from exc

        def shade_cell(cell, fill: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

        def set_cell_width(cell, width_cm: float) -> None:
            cell.width = Cm(width_cm)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width_cm * 567)))
            tc_w.set(qn("w:type"), "dxa")

        def add_label(paragraph, label: str) -> None:
            run = paragraph.add_run(label)
            run.bold = True
            run.font.size = Pt(9)

        def split_sentences(text: str) -> list[str]:
            normalized = " ".join(text.split())
            if not normalized:
                return []
            parts = re.split(r"(?<=[.!?])\s+", normalized)
            return [part.strip() for part in parts if part.strip()]

        def distribute_text(parts: list[str], count: int) -> list[str]:
            if count <= 0:
                return []
            if not parts:
                return ["" for _ in range(count)]
            buckets: list[list[str]] = [[] for _ in range(count)]
            for idx, part in enumerate(parts):
                bucket_idx = min(int(idx * count / max(len(parts), 1)), count - 1)
                buckets[bucket_idx].append(part)
            return [" ".join(bucket).strip() for bucket in buckets]

        def manual_segments(scene: Scene, language: str, image_count: int) -> list[tuple[str, str]]:
            text = scene.texts.get(language, TextPanel())
            slide_panels = scene.slide_panels.get(language, []) if scene.slide_panels else []
            if len(slide_panels) >= image_count:
                panel_segments = [
                    ((panel.body or "").strip(), (panel.speaker_notes or "").strip())
                    for panel in slide_panels[:image_count]
                ]
                if any(body or speaker for body, speaker in panel_segments):
                    return panel_segments
            body_segments = distribute_text(split_sentences(text.body), image_count)
            speaker_segments = distribute_text(split_sentences(text.speaker_notes or text.body), image_count)
            return list(zip(body_segments, speaker_segments))

        def prepare_docx_image(image_path: Path, language: str, scene_idx: int, img_idx: int) -> Path:
            try:
                from PIL import Image as PILImage
            except ImportError as exc:
                raise RuntimeError("Pillow ist fuer die Handbuch-Bildvalidierung nicht installiert.") from exc

            try:
                with PILImage.open(image_path) as img:
                    img.verify()
            except Exception as exc:
                raise RuntimeError(f"Frame ist kein gueltiges Bild fuer das Handbuch: {image_path.name}") from exc

            try:
                with PILImage.open(image_path) as img:
                    img = img.convert("RGB")
                    tmp_dir = settings.workspace_root / "tmp" / "manual-docx-images" / storyboard.video_id / language
                    tmp_dir.mkdir(parents=True, exist_ok=True)
                    safe_path = tmp_dir / f"scene_{scene_idx:03d}_image_{img_idx:03d}.jpg"
                    img.save(safe_path, "JPEG", quality=95, optimize=True)
                    return safe_path
            except Exception as exc:
                raise RuntimeError(f"Frame konnte nicht fuer Word vorbereitet werden: {image_path.name}") from exc

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(210)
        section.page_height = Mm(148)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        styles = doc.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(10)

        # --- Deckblatt ---
        title = doc.add_heading("Clip2Guide-Handbuch", level=0)
        title.runs[0].font.size = Pt(28)
        title.paragraph_format.space_after = Pt(6)

        source_name = Path(storyboard.source_video).stem if storyboard.source_video else ""
        if source_name:
            sub = doc.add_paragraph(source_name)
            sub.runs[0].font.size = Pt(14)
            sub.runs[0].font.color.rgb = None
            sub.paragraph_format.space_after = Pt(20)

        _META_ROWS = [
            ("Sprache", lang),
            ("Erstellt am", date.today().isoformat()),
            ("Anzahl Szenen", str(len(storyboard.scenes))),
            ("Clip2Guide-Version", "0.1.0"),
        ]
        if ai_info:
            _META_ROWS.append(("KI-Modell", ai_info))
        meta_table = doc.add_table(rows=len(_META_ROWS), cols=2)
        meta_table.style = "Table Grid"
        meta_table.autofit = False
        for row, (label, value) in zip(meta_table.rows, _META_ROWS):
            lbl_cell = row.cells[0]
            val_cell = row.cells[1]
            set_cell_width(lbl_cell, 4.0)
            set_cell_width(val_cell, 14.8)
            shade_cell(lbl_cell, "E8EDF2")
            lbl_run = lbl_cell.paragraphs[0].add_run(label)
            lbl_run.bold = True
            lbl_run.font.size = Pt(10)
            val_run = val_cell.paragraphs[0].add_run(value)
            val_run.font.size = Pt(10)

        doc.add_paragraph()

        toc_heading = doc.add_heading("Inhalt", level=2)
        toc_heading.runs[0].font.size = Pt(13)

        toc_table = doc.add_table(rows=1 + len(storyboard.scenes), cols=3)
        toc_table.style = "Table Grid"
        toc_table.autofit = False
        hdr = toc_table.rows[0]
        for cell, text, w in zip(hdr.cells, ["#", "Szene", "Bilder"], [1.2, 15.6, 2.0]):
            set_cell_width(cell, w)
            shade_cell(cell, "2E4057")
            r = cell.paragraphs[0].add_run(text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = None
        for s_idx, scene in enumerate(storyboard.scenes, 1):
            row = toc_table.rows[s_idx]
            text = scene.texts.get(lang, TextPanel())
            scene_title = text.heading or scene.scene_id
            img_count = str(len(scene.image_group) if scene.image_group else 1)
            shade = "F7F9FB" if s_idx % 2 == 0 else "FFFFFF"
            for cell, val, w in zip(row.cells, [str(s_idx), scene_title, img_count], [1.2, 15.6, 2.0]):
                set_cell_width(cell, w)
                shade_cell(cell, shade)
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(10)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        frames_dir = settings.frames_dir / storyboard.video_id
        for scene_idx, scene in enumerate(storyboard.scenes, 1):
            if scene_idx > 1:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            text = scene.texts.get(lang, TextPanel())
            doc.add_heading(f"{scene_idx}. {text.heading or scene.scene_id}", level=1)
            image_group = scene.image_group or ([scene.start_frame] if scene.start_frame else [])
            segments = manual_segments(scene, lang, max(len(image_group), 1))

            for img_idx, filename in enumerate(image_group or [""]):
                if img_idx > 0:
                    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                # Breite des nutzbaren Bereichs: A5 quer 210mm - 2cm Ränder = 188mm
                PAGE_W_CM = 18.8

                table = doc.add_table(rows=2, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.style = "Table Grid"

                img_cell = table.rows[0].cells[0]
                panel_cell = table.rows[1].cells[0]

                for cell in (img_cell, panel_cell):
                    set_cell_width(cell, PAGE_W_CM)

                img_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                panel_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                shade_cell(img_cell, "F3F6FA")
                shade_cell(panel_cell, "FFFFFF")

                img_path = frames_dir / filename
                if filename and img_path.exists():
                    safe_img_path = prepare_docx_image(img_path, lang, scene_idx, img_idx + 1)
                    try:
                        from PIL import Image as _PILImage
                        with _PILImage.open(safe_img_path) as _pil:
                            orig_w, orig_h = _pil.size
                    except Exception:
                        orig_w, orig_h = 1, 1

                    max_w_cm = PAGE_W_CM - 0.4
                    max_h_cm = 5.0
                    if orig_w / max(orig_h, 1) >= max_w_cm / max_h_cm:
                        img_w = Cm(max_w_cm)
                        img_h = None
                    else:
                        img_h = Cm(max_h_cm)
                        img_w = None

                    add_label(img_cell.paragraphs[0], f"Abbildung {scene_idx}.{img_idx + 1}")
                    img_cell.paragraphs[0].add_run("\n")
                    run = img_cell.paragraphs[0].add_run()
                    if img_w:
                        run.add_picture(str(safe_img_path), width=img_w)
                    else:
                        run.add_picture(str(safe_img_path), height=img_h)
                else:
                    img_cell.text = filename or "(kein Bild)"

                body_text, _ = segments[img_idx] if img_idx < len(segments) else ("", "")
                effective_text = body_text or text.body or text.speaker_notes or ""
                p = panel_cell.paragraphs[0]
                p.add_run(effective_text)
                doc.add_paragraph()

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
